import json
import os
import time
import pandas as pd
from docx import Document
from docx.shared import Pt
import google.generativeai as genai
from openai import OpenAI
from dotenv import load_dotenv

# --- Configuration ---
load_dotenv()
INPUT_FILE = 'chat_history_last_month.json'
RECALLS_EXCEL = 'AMC_Recalls.xlsx'
NOTES_WORD = 'AMC_Notes.docx'
IMAGE_DIR = 'downloads'
BATCH_SIZE = 15

# 1. Gemini (Used ONLY for note refinement)
GEMINI_API_KEY = os.getenv('GEMINI_API_KEY')
genai.configure(api_key=GEMINI_API_KEY)
gemini_model = genai.GenerativeModel('gemini-flash-latest')

# 2. DeepSeek (Primary Engine for ALL recall analysis)
DEEPSEEK_API_KEY = os.getenv('DEEPSEEK_API_KEY')
ds_client = OpenAI(api_key=DEEPSEEK_API_KEY, base_url="https://api.deepseek.com")

# Helper to track if Gemini is still "healthy" in this session (only used for notes)
gemini_session_active = True

# --- System Prompt ---
SYSTEM_PROMPT = """
You are an expert AMC (Australian Medical Council) exam preparation assistant.
Your task is to analyze medical messages and classify them into 'Recall' or 'Note'.

1. **Recall**: A medical exam question.
   - Identify the correct answer using the latest evidence-based clinical guidelines.
   - For the explanation field, you MUST provide ALL of the following sections:

     ### ✅ Correct Answer: [Option Letter]. [Option Text]
     Provide a thorough, step-by-step clinical reasoning for WHY this is correct. Include pathophysiology, key diagnostic features, and management pearls relevant to the AMC exam.

     ### ❌ Why the Other Options Are Wrong
     For EACH incorrect option, provide a dedicated sub-bullet explaining exactly why it is excluded. Use this format:
     - **[Option Letter]. [Option Text]**: [Clear clinical reason for exclusion, e.g. different symptoms, different mechanism, different timeline, etc.]

     ### 📚 Relevant Guidelines
     Provide 1–3 relevant, authoritative clinical guidelines or references for this topic. For each, include:
     - **[Guideline Name / Organisation]**: [Brief description of what this guideline covers] — [Full URL link]
     Prefer Australian guidelines (e.g., RACGP, Therapeutic Guidelines Australia, ATAGI, eTG, Choosing Wisely Australia) where applicable, otherwise use international ones (e.g., NICE, UpToDate, AHA, WHO).

   - Output fields: id, date, stem, options (list), answer, explanation, image_path, source_group.
   - Copy the `source_group` value exactly from the input message.

2. **Note**: A clinical knowledge point, guideline summary, or high-yield tip.
   - You MUST categorize each Note using one of these standard categories:
     [Adult General, Cardiology, Dermatology, Endocrinology, Gastroenterology, Haematology, Infectious Diseases, Neurology, Obstetrics, Gynaecology, Paediatric, Psychiatry, Respiratory, Surgery, Emergency, Public Health].
   - Output fields: id, date, category, title, content, image_path, source_group.
   - Copy the `source_group` value exactly from the input message.

Output ONLY a JSON object with two lists: 'recalls' and 'notes'. Ignore ads and noise.
"""

NOTE_REFINEMENT_PROMPT = """
You are a professional medical editor at a top-tier medical journal.
Your task is to take a collection of raw clinical notes (categorized as {category}) and transform them into a highly readable, professionally formatted summary for a Word document.

Instructions:
1. **Consolidate**: Merge overlapping or similar points.
2. **Structure**: Use bullet points, bolded headings, and clear paragraphs.
3. **High-Yield**: Focus on the 'must-know' information for the AMC exam. 
4. **Tone**: Educational, authoritative, and concise.

Raw Notes:
{notes_text}

Output the refined content in plain text with clear Markdown-style headers (e.g., #, ##, ###) and bolding (**text**) so I can parse it into a Word document.
"""

def clean_json_text(text):
    text = text.strip()
    if text.startswith("```json"):
        text = text[7:-3]
    elif text.startswith("```"):
        text = text[3:-3]
    return text.strip()

def process_messages_with_deepseek(messages):
    """Primary engine for all recall + note classification. Uses DeepSeek exclusively."""
    prompt_content = f"Analyze these messages and follow the system instructions carefully. Be thorough and detailed, especially for Recalls:\n\n{json.dumps(messages, ensure_ascii=False)}"
    
    try:
        response = ds_client.chat.completions.create(
            model="deepseek-chat",
            messages=[
                {"role": "system", "content": SYSTEM_PROMPT},
                {"role": "user", "content": prompt_content},
            ],
            stream=False,
            timeout=180  # Extended timeout for detailed responses
        )
        text = clean_json_text(response.choices[0].message.content)
        return json.loads(text)
    except Exception as e:
        print(f"DeepSeek failed: {e}")
        return {"recalls": [], "notes": []}

def process_messages(messages):
    """Routes all message analysis to DeepSeek (primary & only engine for recalls)."""
    print("Processing with DeepSeek...")
    return process_messages_with_deepseek(messages)

def refine_notes_with_ai(category, notes):
    """Refines a list of notes for a specific category using AI."""
    notes_text = ""
    for n in notes:
        notes_text += f"- Title: {n.get('title')}\n  Content: {n.get('content')}\n\n"
        
    prompt = NOTE_REFINEMENT_PROMPT.format(category=category, notes_text=notes_text)
    
    try:
        # We'll use the same engine choice logic here
        global gemini_session_active
        if gemini_session_active:
            try:
                response = gemini_model.generate_content(prompt)
                return response.text.strip()
            except Exception as e:
                if "429" in str(e) or "quota" in str(e).lower():
                    print("Gemini quota exhausted during refinement. Switching to DeepSeek.")
                    gemini_session_active = False
                raise e
        
        # Fallback for refinement too
        response = ds_client.chat.completions.create(
            model="deepseek-chat",
            messages=[{"role": "user", "content": prompt}]
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        print(f"Refinement failed for category {category}: {e}")
        return notes_text # Return raw if refinement fails

def save_to_excel(recalls):
    if not recalls:
        return
    df = pd.DataFrame(recalls)
    # Ensure all columns exist (source_group added to track origin)
    for col in ['id', 'date', 'stem', 'options', 'answer', 'explanation', 'image_path', 'source_group']:
        if col not in df.columns:
            df[col] = None
    
    # Format options list to string
    df['options'] = df['options'].apply(lambda x: "\n".join(x) if isinstance(x, list) else x)
    
    # Reorder so source_group is visible early
    col_order = ['id', 'date', 'source_group', 'stem', 'options', 'answer', 'explanation', 'image_path']
    df = df[col_order]
    
    df.to_excel(RECALLS_EXCEL, index=False)
    print(f"Saved {len(recalls)} recalls to {RECALLS_EXCEL}")

def save_to_word(notes_summary_dict):
    """Saves AI-refined summaries to a Word document."""
    if not notes_summary_dict:
        return
    doc = Document()
    doc.add_heading('AMC Clinical Knowledge Summary', 0)
    
    for category in sorted(notes_summary_dict.keys()):
        doc.add_heading(category, level=1)
        content = notes_summary_dict[category]
        
        # Simple Markdown-to-Word parsing
        lines = content.split('\n')
        for line in lines:
            line = line.strip()
            if not line: continue
            
            if line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            else:
                p = doc.add_paragraph()
                # Handle bold **text**
                parts = line.split('**')
                for i, part in enumerate(parts):
                    run = p.add_run(part)
                    if i % 2 == 1:
                        run.bold = True
                        
    doc.save(NOTES_WORD)
    print(f"Professional notes saved to {NOTES_WORD}")

def main():
    if not os.path.exists(INPUT_FILE):
        print(f"Input file {INPUT_FILE} not found.")
        return
        
    with open(INPUT_FILE, 'r', encoding='utf-8') as f:
        all_data = json.load(f)
        
    print(f"Total messages to analyze: {len(all_data)}")
    
    all_recalls = []
    all_notes = []
    
    for i in range(0, len(all_data), BATCH_SIZE):
        batch = all_data[i : i + BATCH_SIZE]
        print(f"Processing batch {i//BATCH_SIZE + 1} / {(len(all_data)-1)//BATCH_SIZE + 1}...")
        
        # Prepare lean data for AI (only text and minimal metadata)
        lean_batch = []
        for m in batch:
            lean_batch.append({
                "id": m.get('id'),
                "date": m.get('date'),
                "text": m.get('text'),
                "image_path": m.get('image_path'),
                "source_group": m.get('source_group', 'Unknown'),
            })
            
            
        result = process_messages(lean_batch)
        all_recalls.extend(result.get('recalls', []))
        all_notes.extend(result.get('notes', []))
        
        # Intermediate Save for recalls
        save_to_excel(all_recalls)
        
        time.sleep(2) # Protect rate limit

    print("\n--- Refining Notes ---")
    # Group by category
    notes_by_cat = {}
    for n in all_notes:
        cat = n.get('category', 'Uncategorized')
        if cat not in notes_by_cat: notes_by_cat[cat] = []
        notes_by_cat[cat].append(n)
    
    refined_summaries = {}
    for cat, notes_list in notes_by_cat.items():
        print(f"Refining category: {cat}...")
        refined_summaries[cat] = refine_notes_with_ai(cat, notes_list)
        
    save_to_word(refined_summaries)

    print("\n--- Processing Complete ---")
    print(f"Final Count: {len(all_recalls)} Recalls, {len(all_notes)} Notes refined.")

if __name__ == '__main__':
    main()
